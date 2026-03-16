import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import anthropic
import pandas as pd
import json  # 

st.set_page_config(page_title="GMP DocWriter | XI", page_icon="📋", layout="wide", initial_sidebar_state="expanded")

if "lang" not in st.session_state:
    st.session_state.lang = "EN"

def t(en, zh):
    return zh if st.session_state.lang == "ZH" else en

DEFAULTS = {
    "step": 1, "doc_type": None, "basic": {}, "event": {}, "impact": {},
    "rca": {}, "capa": {}, "risk": {}, "chat_history": [],
    "ai_capa_proposal": "", "report_generated": False,
    "step3_gaps": "", "step4_gaps": "", "ra_matrix_result": ""
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v if not isinstance(v, (dict, list)) else (v.copy() if isinstance(v, dict) else [])

@st.cache_resource
def get_client():
    return anthropic.Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])

def ask_claude(system_prompt, user_message):
    msg = get_client().messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=1024,
        system=system_prompt,
        messages=[{"role": "user", "content": user_message}]
    )
    return msg.content[0].text

def lang_prefix():
    return "Reply in Traditional Chinese." if st.session_state.lang == "ZH" else "Reply in English."

def show_gaps(gaps_text):
    for line in gaps_text.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("❌"):
            st.error(line)
        elif line.startswith("⚠️"):
            st.warning(line)
        else:
            st.markdown(line)

SYS_FLAG = (
    "You are a strict GMP technical writing auditor for radiopharmaceutical manufacturing.\n"
    "Your job is to review deviation report text and flag ONLY what is missing or unclear.\n"
    "DO NOT rewrite or improve the text. DO NOT invent or assume any data.\n"
    "Output a bullet list of gaps only, using these prefixes:\n"
    "❌ Missing — something required is completely absent\n"
    "⚠️ Vague — something is present but too imprecise for audit\n"
    "Each bullet must be specific: say exactly what is missing and why it matters for GMP compliance.\n"
    "If something looks complete, do not comment on it."
)

SYS_IMPACT = (
    "You are a strict GMP impact assessment auditor for radiopharmaceutical manufacturing.\n"
    "Review the impact assessment and challenge vague or unsupported claims.\n"
    "Requirements for a compliant impact assessment:\n"
    "- Process impact: specific process steps affected, what was interrupted or changed\n"
    "- Quality impact: specific test data or objective evidence, not assumptions\n"
    "- Patient safety: clear traceability from process change to quality to patient outcome\n"
    "DO NOT rewrite. Flag gaps only using:\n"
    "❌ Missing — required element is absent\n"
    "⚠️ Vague — present but lacks specificity or evidence"
)

SYS_RCA = (
    "You are a GMP root cause analysis expert. Guide the user through 6M analysis.\n"
    "Probe each M - do not accept not applicable without evidence.\n"
    "If user says Man error, probe deeper: training gap? procedure gap? verification gap? workload?\n"
    "Be concise. Ask one focused question at a time."
)

SYS_RISK = (
    "You are a GMP risk assessment expert. Challenge vague impact statements and underestimated risk scores.\n"
    "Ask for specific evidence to justify probability and severity.\n"
    "Point out if quality and patient safety are being treated as independent when they are linked.\n"
    "Be concise and direct."
)

SYS_CAPA = (
    "You are a GMP CAPA expert. Based on the RCA conversation, propose:\n"
    "1. Corrective Action (CA) - fix THIS specific event\n"
    "2. Preventive Action (PA) - address root cause to prevent recurrence (must differ from CA)\n"
    "3. Effectiveness Check - measurable and time-bound verification method\n"
    "Each action must specify owner type (QA/Production/Training). Be specific, not generic."
)

SYS_RA = (
    "You are a GMP risk assessment expert for radiopharmaceutical manufacturing.\n"
    "Based on the deviation details, recommend a risk rating.\n"
    "Output format:\n"
    "1. Recommended Probability (Frequent 5 / Probable 4 / Occasional 3 / Remote 2 / Unlikely 1) + justification\n"
    "2. Recommended Severity (Catastrophic 4 / Critical 3 / Marginal 2 / Negligible 1) + justification\n"
    "3. Risk Score (P x S) and Risk Level (High>=12 / Medium 6-11 / Low<=5)\n"
    "4. Key evidence used. Be specific. Do not accept vague justifications."
)

with st.sidebar:
    col1, col2 = st.columns(2)
    with col1:
        if st.button("English", use_container_width=True,
                     type="primary" if st.session_state.lang == "EN" else "secondary", key="lang_en"):
            st.session_state.lang = "EN"
    with col2:
        if st.button("中文", use_container_width=True,
                     type="primary" if st.session_state.lang == "ZH" else "secondary", key="lang_zh"):
            st.session_state.lang = "ZH"
    st.divider()
    st.markdown(f"### {t('Progress', '填寫進度')}")
    steps_labels = [
        t("Doc Type", "文件類型"), t("Basic Info", "基本資訊"),
        t("Event Detail", "事件描述"), t("Impact", "影響評估"),
        t("Root Cause", "根本原因"), t("CAPA & Risk", "CAPA 與風險")
    ]
    for i, s in enumerate(steps_labels, 1):
        icon = "✅" if st.session_state.step > i else ("🔵" if st.session_state.step == i else "⚪")
        st.markdown(f"{icon} **Step {i}:** {s}")

    st.divider()
    st.markdown(f"### {t('Save / Load', '存檔 / 載入')}")

    # 存檔
    save_data = {
        "doc_type": st.session_state.doc_type,
        "step": st.session_state.step,
        "basic": st.session_state.basic,
        "event": st.session_state.event,
        "impact": st.session_state.impact,
        "rca": st.session_state.rca,
        "capa": st.session_state.capa,
        "risk": st.session_state.risk,
        "ai_capa_proposal": st.session_state.ai_capa_proposal,
        "chat_history": st.session_state.chat_history,
        "ra_matrix_result": st.session_state.get("ra_matrix_result", ""),
    }
    fn = st.session_state.basic.get("dev_number", "draft") + "_save.json"
    st.sidebar.download_button(
        label=t("💾 Save Progress", "💾 儲存進度"),
        data=json.dumps(save_data, ensure_ascii=False, indent=2),
        file_name=fn,
        mime="application/json",
        use_container_width=True
    )

    # 載入
    uploaded = st.sidebar.file_uploader(
        t("📂 Load Progress", "📂 載入存檔"),
        type="json", key="load_file"
    )
    if uploaded is not None:
        loaded = json.load(uploaded)
        for k, v in loaded.items():
            st.session_state[k] = v
        st.rerun()

    st.divider()
    if st.button(t("🔄 Start Over", "🔄 重新開始"), use_container_width=True):
        for k, v in DEFAULTS.items():
            st.session_state[k] = v if not isinstance(v, (dict, list)) else (v.copy() if isinstance(v, dict) else [])
        st.rerun()

st.markdown(f"# 📋 {t('GMP DocWriter', 'GMP 文件撰寫助手')}")
st.markdown(f"*{t('XI Technical Writing Assistant — Deviation Module', 'XI 技術文件撰寫助手 — 偏差模組')}*")
st.divider()

if st.session_state.step == 1:
    st.subheader(f"Step 1 / 6 — {t('Select Document Type', '選擇文件類型')}")
    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button("📄 DEV  " + t("Deviation Report", "偏差報告"), use_container_width=True, key="btn_dev"):
            st.session_state.doc_type = "DEV"; st.session_state.step = 2; st.rerun()
    with c2:
        if st.button("🔬 INV  " + t("Investigation Report", "調查報告"), use_container_width=True, key="btn_inv"):
            st.session_state.doc_type = "INV"; st.session_state.step = 2; st.rerun()
    with c3:
        if st.button("⚡ PRDI  " + t("Pre-Release Deviation", "放行前偏差調查"), use_container_width=True, key="btn_prdi"):
            st.session_state.doc_type = "PRDI"; st.session_state.step = 2; st.rerun()

elif st.session_state.step == 2:
    st.subheader(f"Step 2 / 6 — {t('Basic Information', '基本資訊')} [{st.session_state.doc_type}]")
    b = st.session_state.basic
    c1, c2 = st.columns(2)
    with c1:
        b["dev_number"] = st.text_input(t("Document Number (e.g. DEV-2026-004)", "文件編號"), value=b.get("dev_number", ""))
        b["product"] = st.text_input(t("Product / Radiopharmaceutical", "產品"), value=b.get("product", ""))
        b["lot_number"] = st.text_input(t("Lot / Batch Number", "批號"), value=b.get("lot_number", ""))
        site_opts = ["XWH", "TCM", "TTH", "REN", "Other"]
        b["site"] = st.selectbox(t("Site", "地點"), site_opts, index=site_opts.index(b.get("site", "XWH")))
    with c2:
        b["date_occurrence"] = st.text_input(t("Date of Occurrence (yyyy-Mon-dd)", "發生日期"), value=b.get("date_occurrence", ""))
        b["date_discovery"] = st.text_input(t("Date of Discovery (yyyy-Mon-dd)", "發現日期"), value=b.get("date_discovery", ""))
        b["operator"] = st.text_input(t("Person Involved", "相關人員"), value=b.get("operator", ""))
        area_opts = ["Clinic", "Imaging", "Chemistry", "Other"]
        b["business_area"] = st.selectbox(t("Business Area", "業務範疇"), area_opts, index=area_opts.index(b.get("business_area", "Chemistry")))
    b["source_doc"] = st.text_input(t("Source Document (number + version, e.g. MBR-AV45 v6.0)", "來源文件（需包含編號和版本）"), value=b.get("source_doc", ""))
    etype_opts = ["Non-conformance", "Planned Deviation", "OOS", "OOT", "Continuous Improvement"]
    b["event_type"] = st.selectbox(t("Type of Event", "事件類型"), etype_opts, index=etype_opts.index(b.get("event_type", "Non-conformance")))
    b["title"] = st.text_input(t("Document Title", "文件標題"), value=b.get("title", ""))
    st.session_state.basic = b
    if st.button(t("Next →", "下一步 →"), type="primary", use_container_width=True):
        if not b.get("dev_number") or not b.get("product") or not b.get("title"):
            st.warning(t("⚠️ Some basic fields are empty — you can still proceed, but the report may be incomplete.",
                         "⚠️ 部分基本資訊尚未填寫，仍可繼續，但報告可能不完整。"))
        st.session_state.step = 3
        st.rerun()

elif st.session_state.step == 3:
    st.subheader(f"Step 3 / 6 — {t('Event Description', '事件描述')}")
    e = st.session_state.event
    st.info(t(
        "📌 Every sentence must be independently verifiable. Include exact document references (number+version), specific values, names, exact timestamps.",
        "📌 每句話必須可以獨立驗證。包含精確的文件引用（編號+版本）、具體數值、相關人員姓名、精確時間戳。"
    ))
    e["description"] = st.text_area(t("Description of the Deviation (What/Where/When/Who/Extent)", "偏差描述"), value=e.get("description", ""), height=180)
    e["immediate_action"] = st.text_area(t("Immediate Actions Taken (with timestamp and person responsible)", "已採取的立即措施"), value=e.get("immediate_action", ""), height=100)
    e["extent"] = st.text_area(t("Scope / Extent of Impact", "影響範圍"), value=e.get("extent", ""), height=80)

    if st.button(t("🔍 AI: Flag Writing Gaps", "🔍 AI：找出缺漏"), key="flag_gaps"):
        all_text = (e.get("description", "") + " " + e.get("immediate_action", "") + " " + e.get("extent", "")).strip()
        if all_text:
            with st.spinner(t("Reviewing...", "審查中...")):
                prompt = lang_prefix() + "\n\nReview this deviation report text and flag all gaps:\n\n" + all_text
                st.session_state["step3_gaps"] = ask_claude(SYS_FLAG, prompt)
        else:
            st.warning(t("Please write something first.", "請先填寫內容。"))

    if st.session_state.get("step3_gaps"):
        st.markdown("---")
        st.markdown(f"#### 🔍 {t('AI Gap Review', 'AI 缺漏審查')}")
        show_gaps(st.session_state["step3_gaps"])

        if st.button(t("📝 Generate Improvement Template", "📝 生成補寫模板"), key="gen_template"):
            all_text_for_template = (
                st.session_state.event.get("description", "") + " " +
                st.session_state.event.get("immediate_action", "") + " " +
                st.session_state.event.get("extent", "")
            ).strip()
            if all_text_for_template:
                prompt = (
                    lang_prefix() +
                    "\n\nBased on the gaps identified, rewrite ONLY the original text as an improvement template. "
                    "Do NOT add any new sections, CAPA, root cause, or disposition content that was not in the original text. "
                    "Only fill gaps that exist within the original text provided. "
                    "For every missing or vague piece of information within the original text, insert a blank in this format: _____ (insert relevant detail). "
                    "Keep the original sentence structure as close as possible. "
                    "Do NOT invent or assume any data. "
                    "Output only the improved version of the original text, no explanation, no extra sections.\n\n"
                    "Original text:\n" + all_text_for_template +
                    "\n\nGaps identified:\n" + st.session_state["step3_gaps"]
                )

                with st.spinner(t("Generating template...", "生成模板中...")):
                    st.session_state["step3_template"] = ask_claude(SYS_FLAG, prompt)
            else:
                st.warning(t("Please fill in the description fields first.", "請先填寫描述欄位。"))

        if st.session_state.get("step3_template"):
            st.markdown(f"#### 📝 {t('Improvement Template — edit directly below', '補寫模板 — 可直接在下方編輯')}")
            st.caption(t(
                "⚠️ Replace all blanks with verified data from batch records. Do not estimate.",
                "⚠️ 請將所有空白欄位替換為批次記錄中的已驗證數據，勿估算。"
            ))
            edited = st.text_area(
                t("Edit template here — copy back to fields above when done", "在此編輯模板 — 完成後複製回上方欄位"),
                value=st.session_state["step3_template"],
                height=300,
                key="step3_template_edit"
            )
            st.session_state["step3_template"] = edited

        st.markdown("---")
        if st.button(t("✅ Apply to report fields", "✅ 套用至報告欄位"), key="apply_step3", type="primary"):
                raw = st.session_state["step3_template"]
                cleaned = "\n".join(
                    line for line in raw.split("\n") if not line.strip().startswith("#")
                ).strip()
                lines = [p.strip() for p in cleaned.split("\n\n") if p.strip()]
                if len(lines) >= 1:
                    st.session_state.event["description"] = lines[0]
                if len(lines) >= 2:
                    st.session_state.event["immediate_action"] = lines[1]
                if len(lines) >= 3:
                    st.session_state.event["extent"] = lines[2]
                st.success(t("✅ Applied! Scroll up to review.", "✅ 已套用！請向上捲動確認。"))
                st.rerun()


    st.session_state.event = e
    cb, cn = st.columns(2)
    with cb:
        if st.button(t("← Back", "← 上一步"), use_container_width=True):
            st.session_state.step = 2; st.rerun()
    with cn:
        if st.button(t("Next →", "下一步 →"), type="primary", use_container_width=True):
            if not e.get("description"):
                st.warning(t("⚠️ Description is empty — you can still proceed, but the report may be incomplete.",
                             "⚠️ 偏差描述尚未填寫，仍可繼續，但報告可能不完整。"))
            st.session_state.step = 4; st.rerun()

elif st.session_state.step == 4:
    st.subheader(f"Step 4 / 6 — {t('Impact Assessment', '影響評估')}")
    imp = st.session_state.impact
    st.warning(t(
        "⚠️ Address all three levels separately. Each claim must be supported by objective evidence.",
        "⚠️ 請分別說明以下三個層面，每項主張須有客觀依據支持。"
    ))
    imp["process_impact"] = st.text_area(t("1. Process Impact — specific steps affected, what changed", "1. 製程影響 — 受影響的具體步驟"), value=imp.get("process_impact", ""), height=150)
    imp["quality_impact"] = st.text_area(t("2. Product Quality Impact — test data or objective evidence", "2. 產品品質影響 — 測試數據或客觀依據"), value=imp.get("quality_impact", ""), height=150)
    imp["patient_impact"] = st.text_area(t("3. Patient Safety Impact — traceability: process → quality → patient outcome", "3. 病患安全影響 — 可追溯性：製程 → 品質 → 病患結果"), value=imp.get("patient_impact", ""), height=150)
    if st.session_state.doc_type == "PRDI":
        disp_opts = ["Release — all specs met", "Reject — spec failure", "Conditional release with QA justification"]
        imp["disposition"] = st.selectbox(t("Batch Disposition", "批次處置"), disp_opts)
        imp["disposition_rationale"] = st.text_area(t("Disposition Rationale", "處置理由"), value=imp.get("disposition_rationale", ""), height=100)

    if st.button(t("🔍 AI: Flag Impact Gaps", "🔍 AI：找出影響評估缺漏"), key="flag_impact"):
        text = "Process: " + imp.get("process_impact", "") + "\nQuality: " + imp.get("quality_impact", "") + "\nPatient: " + imp.get("patient_impact", "")
        if text.strip():
            with st.spinner(t("Reviewing...", "審查中...")):
                prompt = lang_prefix() + "\n\nReview this impact assessment:\n\n" + text
                st.session_state["step4_gaps"] = ask_claude(SYS_IMPACT, prompt)
        else:
            st.warning(t("Please fill in the impact fields first.", "請先填寫影響評估欄位。"))

    if st.session_state.get("step4_gaps"):
        st.markdown("---")
        st.markdown(f"#### 🔍 {t('AI Impact Review', 'AI 影響評估審查')}")
        show_gaps(st.session_state["step4_gaps"])

        if st.button(t("📝 Generate Improvement Template", "📝 生成補寫模板"), key="gen_impact_template"):
            all_impact_text = (
                "Process: " + st.session_state.impact.get("process_impact", "") + "\n" +
                "Quality: " + st.session_state.impact.get("quality_impact", "") + "\n" +
                "Patient: " + st.session_state.impact.get("patient_impact", "")
            ).strip()
            if all_impact_text:
                prompt = (
                    lang_prefix() +
                    "\n\nBased on the gaps identified, rewrite ONLY the original text as an improvement template. "
                    "Do NOT add any new sections, CAPA, root cause, or disposition content that was not in the original text. "
                    "Only fill gaps that exist within the original text provided. "
                    "For every missing or vague piece of information, insert a blank in this format: _____ (insert relevant detail). "
                    "Keep the original sentence structure as close as possible. "
                    "Do NOT invent or assume any data. "
                    "Output only the improved version of the original text, no explanation, no extra sections.\n\n"
                    "Original text:\n" + all_impact_text +
                    "\n\nGaps identified:\n" + st.session_state["step4_gaps"]
                )
                with st.spinner(t("Generating template...", "生成模板中...")):
                    st.session_state["step4_template"] = ask_claude(SYS_IMPACT, prompt)
            else:
                st.warning(t("Please fill in the impact fields first.", "請先填寫影響評估欄位。"))

        if st.session_state.get("step4_template"):
            st.markdown(f"#### 📝 {t('Improvement Template — edit directly below', '補寫模板 — 可直接在下方編輯')}")
            st.caption(t(
                "⚠️ Replace all blanks with verified data from batch records. Do not estimate.",
                "⚠️ 請將所有空白欄位替換為批次記錄中的已驗證數據，勿估算。"
            ))
            edited = st.text_area(
                t("Edit template here — copy back to fields above when done", "在此編輯模板 — 完成後複製回上方欄位"),
                value=st.session_state["step4_template"],
                height=300,
                key="step4_template_edit"
            )
            st.session_state["step4_template"] = edited

        st.markdown("---")

    st.session_state.impact = imp
    cb, cn = st.columns(2)
    with cb:
        if st.button(t("← Back", "← 上一步"), use_container_width=True):
            st.session_state.step = 3; st.rerun()
    with cn:
        if st.button(t("Next →", "下一步 →"), type="primary", use_container_width=True):
            if not imp.get("process_impact"):
                st.warning(t("⚠️ Impact assessment is empty — you can still proceed, but the report may be incomplete.",
                             "⚠️ 影響評估尚未填寫，仍可繼續，但報告可能不完整。"))
            st.session_state.step = 5; st.rerun()


elif st.session_state.step == 5:
    st.subheader(f"Step 5 / 6 — {t('Root Cause Analysis', '根本原因分析')}")
    rca = st.session_state.rca
    st.info(t(
        "📌 You must EITHER confirm OR explicitly exclude each 6M category with evidence.",
        "📌 您必須對每個 6M 類別提供確認或明確排除的依據。"
    ))
    ms_keys = ["Man", "Method", "Machine", "Materials", "Mother_Nature", "Measurement"]
    ms_labels = ["Man 人員", "Method 方法", "Machine 機器", "Materials 物料", "Mother Nature 環境", "Measurement 測量"]
    status_opts = [
        t("Confirmed root cause", "確認為根本原因"),
        t("Contributing factor", "促成因素"),
        t("Excluded — with evidence", "已排除（附依據）"),
        t("Under investigation", "調查中")
    ]
    for i, (m_key, m_label) in enumerate(zip(ms_keys, ms_labels)):
        has_content = bool(rca.get("m_" + m_key + "_detail", ""))
        with st.expander(("✅ " if has_content else "⚪ ") + m_label, expanded=(i == 0)):
            curr = rca.get("m_" + m_key + "_status", status_opts[2])
            if curr not in status_opts: curr = status_opts[2]
            rca["m_" + m_key + "_status"] = st.selectbox(t("Status", "狀態"), status_opts, key="s_" + m_key, index=status_opts.index(curr))
            rca["m_" + m_key + "_detail"] = st.text_area(t("Evidence / Detail", "依據 / 說明"), value=rca.get("m_" + m_key + "_detail", ""), height=80, key="d_" + m_key)
    rca["root_cause_statement"] = st.text_area(t("Final Root Cause Statement", "最終根本原因陳述"), value=rca.get("root_cause_statement", ""), height=120)
    st.divider()
    st.markdown(f"#### 🤖 {t('AI RCA Coach', 'AI 根本原因分析輔導')}")
    user_msg = st.text_input(t("Describe your root cause or ask AI to probe further:", "描述根本原因或請 AI 深入探討："), key="rca_input")
    if st.button(t("Ask AI", "詢問 AI"), key="ask_rca"):
        if user_msg:
            with st.spinner(t("Thinking...", "思考中...")):
                prompt = lang_prefix() + "\nEvent: " + st.session_state.basic.get("title", "") + "\nDescription: " + st.session_state.event.get("description", "") + "\nUser: " + user_msg
                resp = ask_claude(SYS_RCA, prompt)
                st.session_state.chat_history.append(("user", user_msg))
                st.session_state.chat_history.append(("ai", resp))
    for role, msg in st.session_state.chat_history[-8:]:
        if role == "ai": st.info("🤖 " + msg)
        else: st.chat_message("user").write(msg)
    if len(st.session_state.chat_history) >= 2:
        st.divider()
        if st.button(t("🤖 AI: Propose CAPA based on this RCA", "🤖 AI：根據 RCA 提議 CAPA"), key="propose_capa", type="primary"):
            chat_log = "\n".join([r.upper() + ": " + m for r, m in st.session_state.chat_history])
            with st.spinner(t("Generating CAPA proposal...", "生成 CAPA 建議中...")):
                prompt = lang_prefix() + "\nEvent: " + st.session_state.basic.get("title", "") + "\nRoot Cause: " + rca.get("root_cause_statement", "") + "\nRCA Log:\n" + chat_log
                st.session_state.ai_capa_proposal = ask_claude(SYS_CAPA, prompt)
        if st.session_state.ai_capa_proposal:
            st.success("🤖 **" + t("AI CAPA Proposal — copy to Step 6", "AI CAPA 建議 — 複製到第 6 步") + "**")
            st.markdown(st.session_state.ai_capa_proposal)
    st.session_state.rca = rca
    cb, cn = st.columns(2)
    with cb:
        if st.button(t("← Back", "← 上一步"), use_container_width=True):
            st.session_state.step = 4; st.rerun()
    with cn:
        if st.button(t("Next →", "下一步 →"), type="primary", use_container_width=True):
            if not rca.get("root_cause_statement"):
                st.warning(t("⚠️ Root cause statement is empty — you can still proceed, but the report may be incomplete.",
                             "⚠️ 根本原因陳述尚未填寫，仍可繼續，但報告可能不完整。"))
            st.session_state.step = 6; st.rerun()

elif st.session_state.step == 6:
    st.subheader(f"Step 6 / 6 — {t('CAPA & Risk Assessment', 'CAPA 與風險評估')}")
    capa = st.session_state.capa
    risk = st.session_state.risk
    if st.session_state.ai_capa_proposal:
        with st.expander(t("💡 View AI CAPA Proposal from Step 5", "💡 查看第 5 步的 AI CAPA 建議"), expanded=True):
            st.markdown(st.session_state.ai_capa_proposal)
    st.markdown(f"### {t('Corrective Action (CA)', '糾正措施 (CA)')}")
    capa["ca_description"] = st.text_area(t("CA Description", "CA 描述"), value=capa.get("ca_description", ""), height=80)
    c1, c2 = st.columns(2)
    with c1: capa["ca_owner"] = st.text_input(t("CA Owner / Dept", "CA 負責人"), value=capa.get("ca_owner", ""))
    with c2: capa["ca_due"] = st.text_input(t("CA Due Date (yyyy-Mon-dd)", "CA 完成日期"), value=capa.get("ca_due", ""))
    st.markdown(f"### {t('Preventive Action (PA)', '預防措施 (PA)')}")
    capa["pa_description"] = st.text_area(t("PA Description (must link to root cause)", "PA 描述（必須連結至根本原因）"), value=capa.get("pa_description", ""), height=80)
    c1, c2 = st.columns(2)
    with c1: capa["pa_owner"] = st.text_input(t("PA Owner / Dept", "PA 負責人"), value=capa.get("pa_owner", ""))
    with c2: capa["pa_due"] = st.text_input(t("PA Due Date (yyyy-Mon-dd)", "PA 完成日期"), value=capa.get("pa_due", ""))
    capa["effectiveness_check"] = st.text_area(t("Effectiveness Check Method", "有效性驗證方法"), value=capa.get("effectiveness_check", ""), height=80)
    st.divider()
    st.markdown(f"### {t('Risk Assessment', '風險評估')}")
    risk["hazard"] = st.text_area(t("Identified Hazard / Risk", "已識別的危害"), value=risk.get("hazard", ""), height=60)
    prob_opts = ["Frequent (5)", "Probable (4)", "Occasional (3)", "Remote (2)", "Unlikely (1)"]
    sev_opts = ["Catastrophic (4)", "Critical (3)", "Marginal (2)", "Negligible (1)"]
    c1, c2 = st.columns(2)
    with c1:
        curr_prob = risk.get("probability", "Unlikely (1)")
        if curr_prob not in prob_opts: curr_prob = "Unlikely (1)"
        risk["probability"] = st.selectbox(t("Probability", "發生概率"), prob_opts, index=prob_opts.index(curr_prob))
    with c2:
        curr_sev = risk.get("severity", "Negligible (1)")
        if curr_sev not in sev_opts: curr_sev = "Negligible (1)"
        risk["severity"] = st.selectbox(t("Severity", "嚴重程度"), sev_opts, index=sev_opts.index(curr_sev))
    risk["risk_justification"] = st.text_area(t("Risk Justification (justify BOTH probability AND severity)", "風險理由（需分別說明概率和嚴重程度）"), value=risk.get("risk_justification", ""), height=100)
    col_a, col_b = st.columns(2)
    with col_a:
        if st.button(t("🤖 AI: Challenge My Risk Score", "🤖 AI：挑戰我的風險評分"), key="challenge_risk", use_container_width=True):
            text = "Hazard: " + risk.get("hazard", "") + "\nProb: " + risk.get("probability", "") + "\nSev: " + risk.get("severity", "") + "\nJustification: " + risk.get("risk_justification", "")
            if text.strip():
                with st.spinner(t("Analysing...", "分析中...")):
                    prompt = lang_prefix() + "\nChallenge this risk assessment:\n" + text + "\nEvent: " + st.session_state.basic.get("title", "")
                    st.warning("🤖 " + ask_claude(SYS_RISK, prompt))
            else: st.warning(t("Please fill in the risk fields first.", "請先填寫風險評估欄位。"))
    with col_b:
        if st.button(t("🧮 AI: Generate RA Matrix", "🧮 AI：生成風險矩陣"), key="gen_ra", use_container_width=True):
            imp_info = st.session_state.impact
            ctx = ("Event: " + st.session_state.basic.get("title", "") +
                   "\nProduct: " + st.session_state.basic.get("product", "") +
                   "\nProcess Impact: " + imp_info.get("process_impact", "") +
                   "\nQuality Impact: " + imp_info.get("quality_impact", "") +
                   "\nPatient Impact: " + imp_info.get("patient_impact", "") +
                   "\nHazard: " + risk.get("hazard", "") +
                   "\nUser Probability: " + risk.get("probability", "") +
                   "\nUser Severity: " + risk.get("severity", "") +
                   "\nUser Justification: " + risk.get("risk_justification", ""))
            with st.spinner(t("Generating Risk Matrix...", "生成風險矩陣中...")):
                st.session_state["ra_matrix_result"] = ask_claude(SYS_RA, lang_prefix() + "\n\n" + ctx)
    if st.session_state.get("ra_matrix_result"):
        st.markdown("---")
        st.markdown(f"#### 🧮 {t('AI Risk Matrix Recommendation', 'AI 風險矩陣建議')}")
        prob_map = {5: "Frequent", 4: "Probable", 3: "Occasional", 2: "Remote", 1: "Unlikely"}
        sev_map = {1: "Negligible", 2: "Marginal", 3: "Critical", 4: "Catastrophic"}
        matrix_rows = []
        for p in [5, 4, 3, 2, 1]:
            row = []
            for s in [1, 2, 3, 4]:
                score = p * s
                if score >= 12: row.append("🔴 " + str(score))
                elif score >= 6: row.append("🟡 " + str(score))
                else: row.append("🟢 " + str(score))
            matrix_rows.append(row)
        df = pd.DataFrame(matrix_rows,
            index=[str(p) + " " + prob_map[p] for p in [5, 4, 3, 2, 1]],
            columns=[str(s) + " " + sev_map[s] for s in [1, 2, 3, 4]])
        st.table(df)
        st.info("🤖 " + st.session_state["ra_matrix_result"])
        st.markdown("---")
    st.session_state.capa = capa
    st.session_state.risk = risk
    cb, cg = st.columns(2)
    with cb:
        if st.button(t("← Back", "← 上一步"), use_container_width=True):
            st.session_state.step = 5; st.rerun()
    with cg:
        if st.button(t("📄 Generate Report", "📄 產生報告"), type="primary", use_container_width=True):
            st.session_state.report_generated = True; st.rerun()

if st.session_state.report_generated:
    b = st.session_state.basic; e = st.session_state.event
    imp = st.session_state.impact; rca = st.session_state.rca
    capa = st.session_state.capa; risk = st.session_state.risk
    st.divider()
    st.success(t("✅ Report Complete — Preview below", "✅ 報告完成 — 預覽如下"))
    with st.expander(t("📄 Report Preview", "📄 報告預覽"), expanded=True):
        st.markdown("## " + str(st.session_state.doc_type) + ": " + b.get("dev_number", ""))
        st.markdown("**" + t("Title", "標題") + ":** " + b.get("title", "") + " | **Site:** " + b.get("site", "") + " | **Product:** " + b.get("product", ""))
        st.markdown("**" + t("Source Doc", "來源文件") + ":** " + b.get("source_doc", "") + " | **" + t("Occurrence", "發生") + ":** " + b.get("date_occurrence", ""))
        st.divider()
        st.markdown("### " + t("Description", "偏差描述"))
        st.write(e.get("description", ""))
        st.markdown("**" + t("Immediate Actions", "立即措施") + ":** " + e.get("immediate_action", ""))
        st.divider()
        st.markdown("### " + t("Impact", "影響"))
        st.markdown("**Process:** " + imp.get("process_impact", ""))
        st.markdown("**Quality:** " + imp.get("quality_impact", ""))
        st.markdown("**Patient Safety:** " + imp.get("patient_impact", ""))
        st.divider()
        st.markdown("### RCA")
        for m_key in ["Man", "Method", "Machine", "Materials", "Mother_Nature", "Measurement"]:
            st.markdown("**" + m_key + ":** " + rca.get("m_" + m_key + "_status", "") + " — " + rca.get("m_" + m_key + "_detail", ""))
        st.markdown("**Root Cause:** " + rca.get("root_cause_statement", ""))
        st.divider()
        st.markdown("### CAPA")
        st.markdown("**CA:** " + capa.get("ca_description", "") + " | " + capa.get("ca_owner", "") + " | " + capa.get("ca_due", ""))
        st.markdown("**PA:** " + capa.get("pa_description", "") + " | " + capa.get("pa_owner", "") + " | " + capa.get("pa_due", ""))
        st.markdown("**Effectiveness:** " + capa.get("effectiveness_check", ""))
        st.divider()
        st.markdown("### " + t("Risk", "風險") + ": " + risk.get("probability", "") + " x " + risk.get("severity", ""))
        st.markdown(risk.get("risk_justification", ""))

    def gen_word():
        doc = Document()
        doc.add_heading("Investigation, Corrective Action, Preventive Action Form", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Document #: " + b.get("dev_number", "") + " | Site: " + b.get("site", "") + " | Business Area: " + b.get("business_area", "") + " | Event Type: " + b.get("event_type", ""))
        doc.add_paragraph("Date of Occurrence: " + b.get("date_occurrence", "") + " | Date of Discovery: " + b.get("date_discovery", "") + " | Person Involved: " + b.get("operator", ""))
        doc.add_paragraph("Source Documentation: " + b.get("source_doc", ""))
        doc.add_heading("Title", 2); doc.add_paragraph(b.get("title", ""))
        doc.add_heading("Description of the Deviation", 2)
        doc.add_paragraph(e.get("description", ""))
        doc.add_paragraph("Immediate Actions: " + e.get("immediate_action", ""))
        doc.add_paragraph("Scope: " + e.get("extent", ""))
        doc.add_heading("Impact Assessment", 2)
        doc.add_paragraph("Process Impact: " + imp.get("process_impact", ""))
        doc.add_paragraph("Product Quality Impact: " + imp.get("quality_impact", ""))
        doc.add_paragraph("Patient Safety Impact: " + imp.get("patient_impact", ""))
        doc.add_heading("Root Cause Analysis — 6M", 2)
        for m_key in ["Man", "Method", "Machine", "Materials", "Mother_Nature", "Measurement"]:
            doc.add_paragraph(m_key + ": " + rca.get("m_" + m_key + "_status", "") + " — " + rca.get("m_" + m_key + "_detail", ""))
        doc.add_paragraph("Root Cause: " + rca.get("root_cause_statement", ""))
        doc.add_heading("CAPA", 2)
        tbl = doc.add_table(rows=3, cols=3); tbl.style = "Table Grid"
        tbl.rows[0].cells[0].text = "Type"; tbl.rows[0].cells[1].text = "Description"; tbl.rows[0].cells[2].text = "Owner / Due"
        tbl.rows[1].cells[0].text = "CA"; tbl.rows[1].cells[1].text = capa.get("ca_description", ""); tbl.rows[1].cells[2].text = capa.get("ca_owner", "") + " / " + capa.get("ca_due", "")
        tbl.rows[2].cells[0].text = "PA"; tbl.rows[2].cells[1].text = capa.get("pa_description", ""); tbl.rows[2].cells[2].text = capa.get("pa_owner", "") + " / " + capa.get("pa_due", "")
        doc.add_paragraph("Effectiveness Check: " + capa.get("effectiveness_check", ""))
        doc.add_heading("Risk Assessment", 2)
        rt = doc.add_table(rows=2, cols=4); rt.style = "Table Grid"
        rt.rows[0].cells[0].text = "Hazard"; rt.rows[0].cells[1].text = "Probability"; rt.rows[0].cells[2].text = "Severity"; rt.rows[0].cells[3].text = "Risk Level"
        try:
            ps = int(risk.get("probability", "Unlikely (1)")[-2])
            ss = int(risk.get("severity", "Negligible (1)")[-2])
            score = ps * ss
            level = "High" if score >= 12 else ("Medium" if score >= 6 else "Low")
        except Exception: score = "N/A"; level = "N/A"
        rt.rows[1].cells[0].text = risk.get("hazard", ""); rt.rows[1].cells[1].text = risk.get("probability", "")
        rt.rows[1].cells[2].text = risk.get("severity", ""); rt.rows[1].cells[3].text = str(level) + " (" + str(score) + ")"
        doc.add_paragraph("Justification: " + risk.get("risk_justification", ""))
        if st.session_state.get("ra_matrix_result"):
            doc.add_heading("AI Risk Matrix Recommendation", 2)
            doc.add_paragraph(st.session_state["ra_matrix_result"])
        doc.add_heading("Approval", 2)
        at = doc.add_table(rows=4, cols=2); at.style = "Table Grid"
        at.rows[0].cells[0].text = "Function"; at.rows[0].cells[1].text = "Signature / Date"
        for i, role in enumerate(["Author/Initiator", "Department Approval", "Quality Assurance Approval"]):
            at.rows[i + 1].cells[0].text = role
        if st.session_state.chat_history:
            doc.add_page_break()
            doc.add_heading("Appendix A — AI-Assisted RCA Conversation Log", 1)
            doc.add_paragraph("Generated: " + datetime.now().strftime("%Y-%b-%d %H:%M") + " | Tool: GMP DocWriter XI")
            for role, msg in st.session_state.chat_history:
                p = doc.add_paragraph()
                label = "Investigator" if role == "user" else "AI RCA Coach"
                run = p.add_run("[" + label + "]: " + msg)
                run.bold = (role == "user"); run.italic = (role == "ai")
                p.paragraph_format.space_after = Pt(6)
        if st.session_state.ai_capa_proposal:
            doc.add_page_break()
            doc.add_heading("Appendix B — AI CAPA Proposal", 1)
            doc.add_paragraph(st.session_state.ai_capa_proposal)
        doc.add_paragraph("\nSOP References: SOP-ISOP-023 | SOP-ISOP-024 | SOP-QA-005")
        buf = BytesIO(); doc.save(buf); buf.seek(0); return buf

    fn = b.get("dev_number", "report") + "_" + b.get("title", "")[:25].replace(" ", "-") + ".docx"
    st.download_button(
        label=t("⬇️ Download Word Report (.docx)", "⬇️ 下載 Word 報告 (.docx)"),
        data=gen_word(), file_name=fn,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary", use_container_width=True
    )
